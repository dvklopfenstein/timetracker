"""Generate a Microsoft Word document containing a table of data"""

__copyright__ = 'Copyright (C) 2025-present, DV Klopfenstein, PhD. All rights reserved.'
__author__ = "DV Klopfenstein, PhD"

##from os import remove
#from os.path import exists
##from os.path import basename
##from os.path import join
##from os.path import abspath
##from os.path import dirname
##from os.path import normpath
##https://python-docx.readthedocs.io/en/latest/
from docx import Document
#from docx.shared import Inches
#from datetime import timedelta
#from datetime import datetime
#from logging import debug
#from csv import reader
#
#from timetracker.utils import orange
#from timetracker.consts import DIRTRK
from timetracker.timetext import get_data_formatted
##from timetracker.cfg.utils import get_username


class WordDoc:
    """Generate a Microsoft Word document containing a table of data"""
    # pylint: disable=too-few-public-methods

    def __init__(self, timedata):
        self.tdata = timedata
        self.nttext = get_data_formatted(timedata)

    def write_doc(self, fout_docx):
        """Write a report into a Microsoft Word document"""
        document = Document()

        document.add_heading('Document Title', 0)

        p = document.add_paragraph('A plain paragraph having some ')
        p.add_run('bold').bold = True
        p.add_run(' and some ')
        p.add_run('italic.').italic = True

        document.add_heading('Heading, level 1', level=1)
        document.add_paragraph('Intense quote', style='Intense Quote')

        document.add_paragraph(
            'first item in unordered list', style='List Bullet'
        )
        document.add_paragraph(
            'first item in ordered list', style='List Number'
        )

        #document.add_picture('monty-truth.png', width=Inches(1.25))

        self._add_table(document)
        document.add_page_break()

        document.save(fout_docx)
        print(f'  WROTE: {fout_docx}')

    def _get_headers(self):
        """Get the number of rows in the timetracking data (self.nttext must have data)"""
        return self.nttext[0]._fields

    def _get_nrows(self):
        """Get the number of rows in the timetracking data (self.nttext must have data)"""
        return len(self.nttext)

    def _get_ncols(self):
        """Get the number of rows in the timetracking data (self.nttext must have data)"""
        return len(self.nttext[0])

    def _add_table(self, doc):
        """Add a table containing timetracking data to a Word document"""
        if not self.nttext:
            return
        table = doc.add_table(rows=1, cols=self._get_ncols())
        for hdr, cell in zip(self._get_headers(), table.rows[0].cells):
            cell.text = hdr
        for ntd in self.nttext:
            row_cells = table.add_row().cells
            for cell, val in zip(row_cells, list(ntd)):
                cell.text = val


# Copyright (C) 2025-present, DV Klopfenstein, PhD. All rights reserved.
